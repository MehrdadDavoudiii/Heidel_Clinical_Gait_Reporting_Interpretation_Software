
from pathlib import Path
import io, json
import numpy as np
import matplotlib.pyplot as plt
import ezc3d
from docx import Document
from docx.shared import Inches

# -------------------- paths & input --------------------
BASE = Path(__file__).parent
ANGLES  = json.loads((BASE / "angles.json").read_text(encoding="utf-8"))
MOMENTS = json.loads((BASE / "moments.json").read_text(encoding="utf-8"))
POWERS  = json.loads((BASE / "powers.json").read_text(encoding="utf-8"))
C3D     = BASE / "Heidel_file.c3d"
STFILE  = BASE / "ST.txt"

if not C3D.exists():
    raise SystemExit(f"Missing C3D: {C3D}")

def by_label(j: dict) -> dict:
    return {str(d.get("label","")).strip(): d for d in j.get("data", [])}

A = by_label(ANGLES)
M = by_label(MOMENTS)
P = by_label(POWERS)

# -------------------- label sets --------------------
KIN_ROWS = [
    ("Pelvis", ("LPelvisAngles","RPelvisAngles")),
    ("Hip",    ("LHipAngles","RHipAngles")),
    ("Knee",   ("LKneeAngles","RKneeAngles")),
    ("Ankle",  ("LAnkleAngles","RAnkleAngles")),
]
MOM_ROWS = [
    ("Hip",   ("LHipMoment","RHipMoment")),
    ("Knee",  ("LKneeMoment","RKneeMoment")),
    ("Ankle", ("LAnkleMoment","RAnkleMoment")),
]
POW_ROW  = [
    ("Hip",   ("LHipPower","RHipPower")),
    ("Knee",  ("LKneePower","RKneePower")),
    ("Ankle", ("LAnklePower","RAnklePower")),
]
PLANES = [
    ("sagittal",  "[FLX–EXT]"),
    ("frontal",   "[ABD–ADD]"),
    ("transverse","[Rotation]"),
]

# -------------------- events (seconds) --------------------
def load_events(c3d_path: Path):
    """Return (left_events, right_events) with items {'type': 'FS'|'FO', 'time': sec} sorted by time."""
    c = ezc3d.c3d(str(c3d_path))
    ev  = c.get("parameters", {}).get("EVENT", {})
    ctx = [str(x).strip() for x in ev.get("CONTEXTS", {}).get("value", [])]
    lab = [str(x).strip() for x in ev.get("LABELS",   {}).get("value", [])]
    tim = np.array(ev.get("TIMES", {}).get("value", []), dtype=float)

    if tim.ndim == 2:  # choose the row that looks like seconds
        row = int(np.nanargmax(np.nanmax(tim, axis=1)))
        sec = tim[row, :]
    else:
        sec = tim

    left, right = [], []
    n = min(len(ctx), len(lab), sec.shape[0] if hasattr(sec,"shape") else len(sec))
    for i in range(n):
        side = ctx[i].lower()
        name = lab[i].lower().replace("  "," ").strip()
        t = float(sec[i])
        if "strike" in name:
            e = {"type":"FS","time":t}
        elif "off" in name:
            e = {"type":"FO","time":t}
        else:
            continue
        (left if side=="left" else right).append(e)
    left.sort(key=lambda e:e["time"]); right.sort(key=lambda e:e["time"])
    return left, right

LEFT_E, RIGHT_E = load_events(C3D)

# -------------------- time-series helpers --------------------
def series_plane(item: dict, plane_key: str):
    t, y = [], []
    for v in item.get("values", []):
        t.append(float(v.get("time", np.nan)))
        val = v.get(plane_key, None)
        y.append(np.nan if val is None else float(val))
    unit = str(item.get("unit","")).strip()
    return np.asarray(t,float), np.asarray(y,float), unit

def cycles_from_FSes(events_side):
    fs = [e["time"] for e in events_side if e["type"]=="FS"]
    return [(fs[i], fs[i+1]) for i in range(len(fs)-1)]

def resample_to_percent(t, y, t0, t1, n=101):
    mask = (t >= t0) & (t <= t1)
    if mask.sum() < 2: return None
    tt, yy = t[mask], y[mask]
    x = (tt - t0) / (t1 - t0)
    grid = np.linspace(0, 1, n)
    good = ~np.isnan(yy)
    if good.sum() < 2: return None
    return np.interp(grid, x[good], yy[good])

def event_percent_in_cycle(t0, t1, events, etype):
    hits = [e["time"] for e in events if e["type"]==etype and t0 <= e["time"] <= t1]
    if not hits: return None
    return 100.0 * (hits[0] - t0) / (t1 - t0)

def mean_sd_and_FO_avg(t, y, same_side):
    cycles = cycles_from_FSes(same_side)
    curves, p_FO = [], []
    for (t0, t1) in cycles:
        cur = resample_to_percent(t, y, t0, t1)
        if cur is None: continue
        curves.append(cur)
        FO = event_percent_in_cycle(t0, t1, same_side, "FO")
        if FO is not None: p_FO.append(FO)
    if not curves:
        return None, None, None
    arr = np.vstack(curves)
    return np.nanmean(arr, axis=0), np.nanstd(arr, axis=0), (np.nanmean(p_FO) if p_FO else None)

def best_cycle_curve(t, y, same_side, *, positive_only=False):
    best = (None, None, None)
    best_peak = -np.inf

    def peak_value(cur):
        curv = cur[np.isfinite(cur)]
        if curv.size == 0:
            return -np.inf
        if positive_only:
            return np.nanmax(curv)
        return np.nanmax(np.abs(curv))

    for (t0, t1) in cycles_from_FSes(same_side):
        cur = resample_to_percent(t, y, t0, t1)
        if cur is None: continue
        pk = peak_value(cur)
        if positive_only and (not np.isfinite(pk) or pk <= 0):
            continue
        if np.isfinite(pk) and pk > best_peak:
            best_peak = pk
            best = (cur, t0, t1)

    if positive_only and best[0] is None:
        return best_cycle_curve(t, y, same_side, positive_only=False)
    return best

# -------------------- plotting utilities --------------------
def auto_ylim(*arrays, pad=0.15):
    vals = []
    for a in arrays:
        if a is None: continue
        vals.extend([np.nanmin(a), np.nanmax(a)])
    if not vals: return None
    lo, hi = float(np.nanmin(vals)), float(np.nanmax(vals))
    span = hi - lo if hi > lo else 1.0
    return lo - pad*span, hi + pad*span

def draw_angles_subplot(ax, title, unit, left_item, right_item, plane_key):
    x = np.linspace(0, 100, 101)
    tL, yL, _ = series_plane(left_item,  plane_key)
    tR, yR, _ = series_plane(right_item, plane_key)

    meanL, sdL, FO_L = mean_sd_and_FO_avg(tL, yL, LEFT_E)
    meanR, sdR, FO_R = mean_sd_and_FO_avg(tR, yR, RIGHT_E)

    if meanL is not None:
        ax.plot(x, meanL, color="red",  linewidth=2.2)
        ax.fill_between(x, meanL - sdL, meanL + sdL, color="red", alpha=0.15)
    if meanR is not None:
        ax.plot(x, meanR, color="blue", linewidth=2.2)
        ax.fill_between(x, meanR - sdR, meanR + sdR, color="blue", alpha=0.15)

    if FO_L is not None: ax.axvline(FO_L, color="red",  linestyle="--", linewidth=1.6, alpha=0.9)
    if FO_R is not None: ax.axvline(FO_R, color="blue", linestyle="--", linewidth=1.6, alpha=0.9)

    ax.set_xlim(0, 100)
    ylims = auto_ylim(
        None if meanL is None else meanL - sdL,
        None if meanL is None else meanL + sdL,
        None if meanR is None else meanR - sdR,
        None if meanR is None else meanR + sdR
    )
    if ylims: ax.set_ylim(*ylims)
    ax.set_title(title, fontsize=11)
    ax.set_xlabel("Gait Cycle (%)")
    ax.set_ylabel(unit)
    ax.grid(True, alpha=0.2, linestyle=":")

def draw_kinetics_bestcycle_subplot(ax, title, unit, left_item, right_item, plane_key, *, positive_only=False, scale=1.0):
    x = np.linspace(0, 100, 101)

    tL, yL, _ = series_plane(left_item, plane_key)
    tR, yR, _ = series_plane(right_item, plane_key)

    # apply scaling
    yL = yL / scale
    yR = yR / scale

    L_curve, L_t0, L_t1 = best_cycle_curve(tL, yL, LEFT_E, positive_only=positive_only)
    R_curve, R_t0, R_t1 = best_cycle_curve(tR, yR, RIGHT_E, positive_only=positive_only)



    if L_curve is not None:
        ax.plot(x, L_curve, color="red", linewidth=2.2)
        L_FO = event_percent_in_cycle(L_t0, L_t1, LEFT_E, "FO") if L_t0 is not None else None
        if L_FO is not None: ax.axvline(L_FO, color="red", linestyle="--", linewidth=1.6, alpha=0.9)
    if R_curve is not None:
        ax.plot(x, R_curve, color="blue", linewidth=2.2)
        R_FO = event_percent_in_cycle(R_t0, R_t1, RIGHT_E, "FO") if R_t0 is not None else None
        if R_FO is not None: ax.axvline(R_FO, color="blue", linestyle="--", linewidth=1.6, alpha=0.9)

    ax.set_xlim(0, 100)
    ylims = auto_ylim(L_curve, R_curve)
    if ylims: ax.set_ylim(*ylims)
    ax.set_title(title, fontsize=11)
    ax.set_xlabel("Gait Cycle (%)")
    ax.set_ylabel(unit)
    ax.grid(True, alpha=0.2, linestyle=":")

def fig_to_doc(doc: Document, fig, width_in=6.5, heading=None):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig); buf.seek(0)
    if heading: doc.add_heading(heading, level=1)
    doc.add_picture(buf, width=Inches(width_in))
    buf.close()

# -------------------- read ST.txt --------------------
def parse_st_file(st_path: Path):
    """Return dict with cadence, speed, foot off, ssup, dsup for left & right"""
    lines = st_path.read_text(errors="ignore").splitlines()
    mean_lines = [ln.strip() for ln in lines if ln.strip().startswith("Mean:")]
    if len(mean_lines) < 2:
        return None
    right_tokens = mean_lines[0].split()[1:]
    left_tokens  = mean_lines[1].split()[1:]
    right_vals = [float(x) for x in right_tokens]
    left_vals  = [float(x) for x in left_tokens]
    return {
        "Right": {
            "Cadence": right_vals[0],
            "Speed":   right_vals[1],
            "FootOff": right_vals[6],
            "SSup":    right_vals[7],
            "DSup":    right_vals[8],
        },
        "Left": {
            "Cadence": left_vals[0],
            "Speed":   left_vals[1],
            "FootOff": left_vals[6],
            "SSup":    left_vals[7],
            "DSup":    left_vals[8],
        }
    }

# -------------------- report assembly --------------------
def build_report():
    doc = Document()

    # ----- Front Page -----
    doc.add_heading("Gait Analysis Report", 0)
    doc.add_paragraph("Patient ID: _______________")
    doc.add_paragraph("Name: ____________________")
    doc.add_paragraph("Address: ____________________")
    doc.add_paragraph("Birth Date: _______________")
    doc.add_paragraph("Examination Date: __________")
    doc.add_paragraph("Height: ______ cm")
    doc.add_paragraph("Weight: ______ kg")
    doc.add_paragraph("Diagnosis: ____________________")

    # Subject parameters (weight, height) from C3D
    c = ezc3d.c3d(str(C3D))
    subj = c["parameters"].get("SUBJECT", {})
    weight = subj.get("WEIGHT", {}).get("value", [None])[0]
    height = subj.get("HEIGHT", {}).get("value", [None])[0]
    if weight: doc.add_paragraph(f"Weight: {weight:.1f} kg")
    if height: doc.add_paragraph(f"Height: {height:.1f} cm")

    # Spatiotemporal from ST.txt
    if STFILE.exists():
        st = parse_st_file(STFILE)
        if st:
            doc.add_heading("Spatiotemporal Parameters", level=1)
            doc.add_paragraph(
                f"Cadence: {st['Right']['Cadence']:.1f} (R), {st['Left']['Cadence']:.1f} (L) steps/min\n"
                f"Speed: {st['Right']['Speed']:.2f} (R), {st['Left']['Speed']:.2f} (L) m/s\n"
                f"Foot Off: {st['Right']['FootOff']:.1f}% (R), {st['Left']['FootOff']:.1f}% (L)\n"
                f"Single Support: {st['Right']['SSup']:.1f}% (R), {st['Left']['SSup']:.1f}% (L)\n"
                f"Double Support: {st['Right']['DSup']:.1f}% (R), {st['Left']['DSup']:.1f}% (L)"
            )

    doc.add_page_break()

    # ----- Page 1: Kinematics
    fig1, axes = plt.subplots(nrows=4, ncols=3, figsize=(10, 12))
    fig1.suptitle("Angles", fontsize=16)
    for r, (row_name, (L, R)) in enumerate(KIN_ROWS):
        if L not in A or R not in A:
            for c in range(3): axes[r, c].axis("off")
            continue
        for c, (plane_key, plane_tag) in enumerate(PLANES):
            ax = axes[r, c]
            title = f"{row_name} {plane_tag}"
            unit  = A[L].get("unit") or A[R].get("unit") or "deg"
            draw_angles_subplot(ax, title, unit, A[L], A[R], plane_key)
    fig1.tight_layout(rect=[0, 0.02, 1, 0.96])
    fig_to_doc(doc, fig1, heading="Angles")

    # ----- Page 2: Moments
    fig2, axes = plt.subplots(nrows=3, ncols=3, figsize=(10, 9))
    fig2.suptitle("Moments", fontsize=16)
    for r, (row_name, (L, R)) in enumerate(MOM_ROWS):
        if L not in M or R not in M:
            for c in range(3): axes[r, c].axis("off")
            continue
        for c, (plane_key, plane_tag) in enumerate(PLANES):
            ax = axes[r, c]
            title = f"{row_name} {plane_tag}"
            unit  = M[L].get("unit") or M[R].get("unit") or "Nm/kg"
            draw_kinetics_bestcycle_subplot(ax, title, unit, M[L], M[R], plane_key, scale=1000.0)

    fig2.tight_layout(rect=[0, 0.02, 1, 0.96])
    fig_to_doc(doc, fig2, heading="Moments")

    # ----- Power
    fig3, axes = plt.subplots(nrows=1, ncols=3, figsize=(10, 3.5))
    fig3.suptitle("Power (Sagittal)", fontsize=16)
    for c, (row_name, (L, R)) in enumerate(POW_ROW):
        ax = axes[c] if isinstance(axes, np.ndarray) else axes
        if L in P and R in P:
            title = f"{row_name} [FLX–EXT]"
            unit  = P[L].get("unit") or P[R].get("unit") or "W/kg"
            draw_kinetics_bestcycle_subplot(ax, title, unit, P[L], P[R], "sagittal", positive_only=True)
        else:
            ax.axis("off")
    fig3.tight_layout(rect=[0, 0.02, 1, 0.96])
    fig_to_doc(doc, fig3, heading="Power")

    # save docx
    doc.save(BASE / "motion_report.docx")

# -------------------- run --------------------
if __name__ == "__main__":
    build_report()
