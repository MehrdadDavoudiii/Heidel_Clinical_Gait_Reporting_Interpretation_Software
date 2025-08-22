
import os, sys, base64, zipfile, mimetypes
from pathlib import Path
from docx import Document
from openai import OpenAI
from docx2pdf import convert

# ------------ Paths ------------
BASE = Path(__file__).parent.resolve()
INPUT_DOCX = BASE / "motion_report.docx"
OUTPUT_DOCX = BASE / "motion_report_with_conclusion.docx"
OUTPUT_PDF  = BASE / "motion_report_with_conclusion.pdf"

MODEL = "gpt-5"

# ------------ API Key ------------
def get_api_key() -> str:
    if len(sys.argv) > 1 and sys.argv[1].strip():
        return sys.argv[1].strip()
    if os.getenv("OPENAI_API_KEY"):
        return os.getenv("OPENAI_API_KEY").strip()
    key = input("Enter OPENAI_API_KEY: ").strip()
    if not key:
        raise SystemExit("No API key provided.")
    return key

# ------------ Helpers ------------
def extract_text(docx_path: Path) -> str:
    doc = Document(str(docx_path))
    return "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())

def extract_images(docx_path: Path):
    imgs = []
    with zipfile.ZipFile(str(docx_path), "r") as z:
        for name in z.namelist():
            if not name.startswith("word/media/"):
                continue
            raw = z.read(name)
            ext = Path(name).suffix
            mime, _ = mimetypes.guess_type(f"x{ext}")
            if not mime:
                mime = "image/png"
            imgs.append((raw, mime))
    return imgs

def to_data_url(mime: str, raw: bytes) -> str:
    return f"data:{mime};base64,{base64.b64encode(raw).decode()}"

def analyze_with_gpt5(client: OpenAI, report_text: str, images):
    content = [
        {"type": "text", "text": (
            "You are given a gait analysis report with text and figures.\n"
            "Interpret kinematics, kinetics (moments), and power from the charts. "
            "Do not say 'no data provided'. Use both text and figures.\n\n"
            "Organize output:\n"
            "1) Spatiotemporal Summary\n"
            "2) Kinematics\n"
            "3) Kinetics (Moments)\n"
            "4) Power\n"
            "5) Key Abnormalities (numbered)\n"
            "6) Clinical Conclusion (short paragraph)"
        )}
    ]
    if report_text:
        content.append({"type": "text", "text": f"Report text:\n{report_text}"})
    for raw, mime in images:
        content.append({"type": "image_url", "image_url": {"url": to_data_url(mime, raw)}})

    resp = client.chat.completions.create(
        model=MODEL,
        messages=[{"role": "user", "content": content}],
    )
    return resp.choices[0].message.content.strip()

def append_interpretation(docx_in: Path, interpretation: str, docx_out: Path):
    doc = Document(str(docx_in))
    doc.add_heading("Interpretation & Clinical Conclusion", level=1)
    for line in interpretation.split("\n"):
        doc.add_paragraph(line.strip())
    doc.save(str(docx_out))

# ------------ Main ------------
if __name__ == "__main__":
    if not INPUT_DOCX.exists():
        raise SystemExit(f"Missing file: {INPUT_DOCX}")

    api_key = get_api_key()
    client = OpenAI(api_key=api_key)

    print("Reading report text…")
    text = extract_text(INPUT_DOCX)

    print("Extracting images…")
    imgs = extract_images(INPUT_DOCX)

    print("Requesting GPT-5 interpretation…")
    interpretation = analyze_with_gpt5(client, text, imgs)

    print("Appending interpretation…")
    append_interpretation(INPUT_DOCX, interpretation, OUTPUT_DOCX)

    print("Exporting PDF…")
    try:
        convert(str(OUTPUT_DOCX), str(OUTPUT_PDF))
        print(f"Done! PDF saved: {OUTPUT_PDF}")
    except Exception as e:
        print(f"Skipped PDF export: {e}")
